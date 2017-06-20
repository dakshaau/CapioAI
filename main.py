import requests, json
import os, sys
from requests.exceptions import HTTPError
from docx import Document
from docx.enum.dml import MSO_THEME_COLOR
from docx.shared import RGBColor

class timestamp(object):
	'''
	This class accomodates the timestamp conversion of time given as seconds.
	'''
	hours = 0
	mins = 0
	secs = 0
	millis = 0
	
	def __init__(self, seconds):
		if seconds >= 86400: # Total number of seconds in a day
			seconds %= 86400
		if seconds < 0:
			return
		self.hours, temp = divmod(seconds, 3600)
		self.mins, temp = divmod(temp, 60)
		self.secs, self.millis = divmod(temp, 1)
		self.millis *= 100

	def __str__(self):
		'''
		Returns a string in the form 'HH:MM:SS.mm'
		'''
		return '{:02.0f}:{:02.0f}:{:02.0f}.{:02.0f}'.format(self.hours, self.mins, self.secs, self.millis)


def getTranscript(transcriptID, APIKey):
	'''
	This function requests the capioai-api for a transcript using the APIKey and TranscriptID
	provided by the user.

	Parameters:
		transcriptID: string
		APIKey: string

	Returns:
		transcript: string, if error occured while handling the request, or
							server rejects the request
									OR
					list of dicts, if server accepts the request
	'''

	try:
		'''
		Making sure that the input arguments are strings
		'''
		assert type(transcriptID) is str
		assert type(APIKey) is str
	except AssertionError:
		return 'Invalid Parameter(s)'


	Error_codes = {'301': 'Moved Permanently',
					'302': 'Found',
					'304': 'Not Modified',
					'400': 'Bad Request',
					'401': 'Unauthorized',
					'403': 'Forbidden',
					'404': 'Not Found',
					'405': 'Method Not Allowed',
					'410': 'Gone',
					'415': 'Unsupported Media Type',
					'422': 'Unprocessable Entry',
					'429': 'Too Many Requests',
					'500': 'Internal Server Error'
					}
	
	url = 'https://api.capio.ai/v1/speech/transcript/{}'.format(transcriptID)
	head = {'apiKey': '{}'.format(APIKey), 'word_confidence': 'true'}
	try:
		result = requests.get(url, headers=head)
	except ConnectionError:
		print('Connection Failed')
		return 'Unable to Connect'
	
	transcript = []
	try:
		result.raise_for_status()
	except HTTPError:
		print('Error Code: {}'.format(result.status_code))
		return Error_codes[str(result.status_code)]
	else:
		transcript = result.json()
		return transcript

def createDocx(results, transcriptID):
	'''
	This method creates a formattted document from the transcript
	retreived from the getTrancript function

	Parameters:
		results: Response received from the CapioAI, list of dicts
		transcriptID: TranscriptID used to retrieve response from CapioAI, string

	Returns:
		doc: docx.Document object if everything is fine, else NoneType
		path: string, containing path to the saved document OR error message
	'''
	purple = RGBColor(0x64, 0x62, 0x96) # Replicating the timestamp color provided in the example file
	red = RGBColor(0xfc, 0x2a, 0x35) # Replicating the 'red' color provided in the example file
	doc = Document()

	try:
		'''
		The following assertions make sure that the results argument parameter has the following format:

		[
			{
				'result': <value>
				'result_index': <int value>
			}
		]
		'''
		assert type(transcriptID) is str
		assert type(results) is list
		assert len(results) > 0
		assert type(results[0]) is dict
		assert 'result_index' in results[0]
		assert 'result' in results[0]
		assert type(results[0]['result_index']) is int
		assert -1 < results[0]['result_index'] < len(results)
	except AssertionError:
		return None, "Improper 'results' format!"
	
	results = sorted(results, key=lambda x: x['result_index']) # Sorting the results in order just in case of some erroneous retreival from the server
	for result in results:
		
		alternatives = result['result']

		try:
			'''
			The following assertions make sure that each 'result' has the following format:

			[
				{
					'alternative': [
									{
										'confidence': <float value>
									}
					]
				}
			]
			'''
			assert type(alternatives) is list
			assert type(alternatives[0]) is dict
			assert 'alternative' in alternatives[0]
			assert type(alternatives[0]['alternative']) is list
			assert type(alternatives[0]['alternative'][0]) is dict
			assert 'confidence' in alternatives[0]['alternative'][0]
			assert (type(alternatives[0]['alternative'][0]['confidence']) is float) or (type(alternatives[0]['alternative'][0]['confidence']) is int)
			assert alternatives[0]['alternative'][0]['confidence'] <= 1
		except AssertionError:
			return None, "Improper 'result' format!"

		alternatives = sorted(alternatives[0]['alternative'], key=lambda x: x['confidence'], reverse=True) # Sorting w.r.t scentence level confidence
		scentence = alternatives[0] # Only taking into account the best transcript
		del alternatives
		
		try:
			'''
			Thw following assertions make sure that 'scentence' has the folowing format:

			{
				'words':[
						{
							'from': <float value>
							'confidence': <float value>
							'word': <string>
						}
				]
			}
			'''
			assert type(scentence) is dict
			assert 'words' in scentence
			assert type(scentence['words']) is list
			assert type(scentence['words'][0]) is dict
			assert 'from' in scentence['words'][0]
			assert 'confidence' in scentence['words'][0]
			assert 'word' in scentence['words'][0]
			assert (type(scentence['words'][0]['from']) is float) or (type(scentence['words'][0]['from']) is int)
			assert scentence['words'][0]['from'] >= 0
			assert (type(scentence['words'][0]['confidence']) is float) or (type(scentence['words'][0]['confidence']) is int)
			assert scentence['words'][0]['confidence'] <= 1
			assert type(scentence['words'][0]['word']) is str
		except AssertionError:
			return None, "Improper 'scentence' format!"

		words = sorted(scentence['words'], key = lambda x: x['from']) # Sorting w.r.t time
		start_time = timestamp(words[0]['from'])

		para = doc.add_paragraph('')
		r = para.add_run('{}\t'.format(start_time))
		r.bold = True
		r.font.color.rgb = purple
		for word in words:
			r = para.add_run(' {}'.format(word['word']))
			if word['confidence'] < 0.75:
				# Marking words with low confidence in red color
				r.font.color.rgb = red

	doc.save('{}.docx'.format(transcriptID))
	path = os.path.abspath('{}.docx'.format(transcriptID))
	return doc, path

if __name__ == '__main__':
	transcriptID = '593f237fbcae700012ba8fcd'
	APIKey = '262ac9a0c9ba4d179aad4c0b9b02120a'
	if len(sys.argv) == 3:
		transcriptID = sys.argv[1]
		APIKey = sys.argv[2]
	else:
		print('Invalid Arguments!')
		print('Usage:')
		print()
		print('python main.py <transcriptID> <APIKey>')
		print()
		exit()
	transcript = getTranscript(transcriptID, APIKey)
	if type(transcript) is str:
		print(transcript)
	elif len(transcript) > 0:
		doc, path = createDocx(transcript, transcriptID)
		if doc is None:
			print(path)
		else:
			print('File saved at: {}'.format(path))
	
