import requests, json
import os, sys
from requests.exceptions import HTTPError
from docx import Document
from docx.enum.dml import MSO_THEME_COLOR
from docx.shared import RGBColor

class timestamp(object):
	'''
	This class accomodates the timestamp conversion of time given as seconds.
	'''
	hours = 0
	mins = 0
	secs = 0
	millis = 0
	
	def __init__(self, seconds):
		if seconds >= 86400: # Total number of seconds in a day
			seconds %= 86400
		if seconds < 0:
			return
		self.hours, temp = divmod(seconds, 3600)
		self.mins, temp = divmod(temp, 60)
		self.secs, self.millis = divmod(temp, 1)
		self.millis *= 100

	def __str__(self):
		'''
		Returns a string in the form 'HH:MM:SS.mm'
		'''
		return '{:02.0f}:{:02.0f}:{:02.0f}.{:02.0f}'.format(self.hours, self.mins, self.secs, self.millis)


def getTranscript(transcriptID, APIKey):
	'''
	This function requests the capioai-api for a transcript using the APIKey and TranscriptID
	provided by the user.

	Parameters:
		transcriptID: string
		APIKey: string

	Returns:
		transcript: string, if error occured while handling the request, or
							server rejects the request
									OR
					list of dicts, if server accepts the request
	'''

	Error_codes = {'301': 'Moved Permanently',
					'302': 'Found',
					'304': 'Not Modified',
					'400': 'Bad Request',
					'401': 'Unauthorized',
					'403': 'Forbidden',
					'404': 'Not Found',
					'405': 'Method Not Allowed',
					'410': 'Gone',
					'415': 'Unsupported Media Type',
					'422': 'Unprocessable Entry',
					'429': 'Too Many Requests',
					'500': 'Internal Server Error'
					}
	
	url = 'https://api.capio.ai/v1/speech/transcript/{}'.format(transcriptID)
	head = {'apiKey': '{}'.format(APIKey), 'word_confidence': 'true'}
	try:
		result = requests.get(url, headers=head)
	except ConnectionError:
		print('Connection Failed')
		return 'Unable to Connect'
	
	transcript = []
	try:
		result.raise_for_status()
	except HTTPError:
		print('Error Code: {}'.format(result.status_code))
		return Error_codes[str(result.status_code)]
	else:
		transcript = result.json()
		return transcript

def createDocx(results, transcriptID):
	'''
	This method creates a formattted document from the transcript
	retreived from the getTrancript function

	Parameters:
		results: Response received from the CapioAI, list of dicts
		transcriptID: TranscriptID used to retrieve response from CapioAI, string

	Returns:
		doc: docx.Document object if everything is fine, else NoneType
		path: string, containing path to the saved document OR error message
	'''
	purple = RGBColor(0x64, 0x62, 0x96) # Replicating the timestamp color provided in the example file
	red = RGBColor(0xfc, 0x2a, 0x35) # Replicating the 'red' color provided in the example file
	doc = Document()

	results = sorted(results, key=lambda x: x['result_index']) # Sorting the results in order just in case of some erroneous retreival from the server
	for result in results:
		
		alternatives = result['result']

		alternatives = sorted(alternatives[0]['alternative'], key=lambda x: x['confidence'], reverse=True) # Sorting w.r.t scentence level confidence
		scentence = alternatives[0] # Only taking into account the best transcript
		del alternatives
		
		words = sorted(scentence['words'], key = lambda x: x['from']) # Sorting w.r.t time
		start_time = timestamp(words[0]['from'])

		para = doc.add_paragraph('')
		r = para.add_run('{}\t'.format(start_time))
		r.bold = True
		r.font.color.rgb = purple
		for word in words:
			r = para.add_run(' {}'.format(word['word']))
			if word['confidence'] < 0.75:
				# Marking words with low confidence in red color
				r.font.color.rgb = red

	doc.save('{}.docx'.format(transcriptID))
	path = os.path.abspath('{}.docx'.format(transcriptID))
	return doc, path

if __name__ == '__main__':
	transcriptID = '593f237fbcae700012ba8fcd'
	APIKey = '262ac9a0c9ba4d179aad4c0b9b02120a'
	if len(sys.argv) == 3:
		transcriptID = sys.argv[1]
		APIKey = sys.argv[2]
	else:
		print('Invalid Arguments!')
		print('Usage:')
		print()
		print('python main.py <transcriptID> <APIKey>')
		print()
		exit()
	transcript = getTranscript(transcriptID, APIKey)
	if type(transcript) is str:
		print(transcript)
	elif len(transcript) > 0:
		doc, path = createDocx(transcript, transcriptID)
		if doc is None:
			print(path)
		else:
			print('File saved at: {}'.format(path))
	
